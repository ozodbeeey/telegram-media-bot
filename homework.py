import wikipedia
import os
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_homework(topic, output_path):
    """
    Generates a Word document (referat) from Wikipedia based on the topic.
    """
    # Try Uzbek first, then fallback to English
    wikipedia.set_lang('uz')
    
    try:
        page = wikipedia.page(topic)
    except wikipedia.exceptions.PageError:
        # Try English if Uzbek fails
        wikipedia.set_lang('en')
        try:
            page = wikipedia.page(topic)
        except wikipedia.exceptions.PageError:
            return False, "Kechirasiz, bu mavzu bo'yicha ma'lumot topilmadi."
        except wikipedia.exceptions.DisambiguationError as e:
            return False, f"Juda ko'p ma'lumot topildi. Aniqroq yozing. Masalan: {e.options[:3]}"
    except wikipedia.exceptions.DisambiguationError as e:
        return False, f"Juda ko'p ma'lumot topildi. Aniqroq yozing. Masalan: {e.options[:3]}"
    except Exception as e:
        return False, f"Xatolik: {e}"

    # Create Document
    doc = Document()
    
    # Title
    title = doc.add_heading(page.title, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph(f"Mavzu: {page.title}")
    doc.add_paragraph(f"Manba: Wikipedia ({page.url})")
    doc.add_paragraph("") # Space

    # Image (if available)
    if page.images:
        img_url = page.images[0]
        # Filter for valid image types usually on wiki (jpg, png)
        # Wiki often returns svgs or page icons which PIL/docx might struggle with.
        valid_extensions = ('.jpg', '.jpeg', '.png')
        for img in page.images:
             if img.lower().endswith(valid_extensions):
                 img_url = img
                 break
        
        try:
            response = requests.get(img_url, stream=True)
            if response.status_code == 200:
                temp_img = "temp_img.jpg"
                with open(temp_img, 'wb') as f:
                    f.write(response.content)
                
                doc.add_picture(temp_img, width=Inches(5))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                if os.path.exists(temp_img):
                    os.remove(temp_img)
        except:
            pass # Ignore image errors

    # Summary
    doc.add_heading('Kirish', level=1)
    doc.add_paragraph(page.summary)

    # Content loop (Simple approach: just get content, but page.content usually has everything)
    # We will try to format it slightly better by splitting by '==', but page.content is a single string.
    # For MVP, dumping content is fine, but let's try to make headings bold.
    
    full_content = page.content
    
    # Split by double newlines to manage paragraphs
    paragraphs = full_content.split('\n\n')
    
    for para in paragraphs:
        clean_para = para.strip()
        if not clean_para:
            continue
            
        # Detect headings (Wiki format often uses == Heading ==)
        if clean_para.startswith('==') and clean_para.endswith('=='):
            heading_text = clean_para.replace('=', '').strip()
            doc.add_heading(heading_text, level=2)
        else:
            doc.add_paragraph(clean_para)

    # Save
    doc.save(output_path)
    return True, "Muvaffaqiyatli saqlandi."
